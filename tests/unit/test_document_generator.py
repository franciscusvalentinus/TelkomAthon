"""Unit tests for DocumentGenerator"""

import pytest
from io import BytesIO
from docx import Document

from src.processors.document_generator import DocumentGenerator
from src.models.entities import SyllabusMaterials, TLO, Performance, ELO


@pytest.fixture
def document_generator():
    """Create a DocumentGenerator instance"""
    return DocumentGenerator()


@pytest.fixture
def sample_materials():
    """Create sample syllabus materials for testing"""
    tlo1 = TLO(
        id="tlo-1",
        org_id="org-1",
        course_type="B2B",
        text="Peserta mampu memahami konsep dasar bisnis B2B",
        is_selected=True
    )
    
    tlo2 = TLO(
        id="tlo-2",
        org_id="org-1",
        course_type="B2B",
        text="Peserta mampu mengimplementasikan strategi penjualan B2B",
        is_selected=True
    )
    
    perf1 = Performance(
        id="perf-1",
        tlo_ids=["tlo-1"],
        text="Mengidentifikasi karakteristik pasar B2B",
        is_selected=True
    )
    
    elo1 = ELO(
        id="elo-1",
        performance_ids=["perf-1"],
        text="Menjelaskan perbedaan antara pasar B2B dan B2C",
        is_selected=True
    )
    
    return SyllabusMaterials(
        organization_summary="PT Example adalah perusahaan teknologi terkemuka",
        organization_context="Perusahaan fokus pada solusi enterprise",
        selected_tlos=[tlo1, tlo2],
        selected_performances=[perf1],
        selected_elos=[elo1],
        course_type="B2B"
    )


def test_create_syllabus_document_returns_bytes(document_generator, sample_materials):
    """Test that create_syllabus_document returns bytes"""
    result = document_generator.create_syllabus_document(sample_materials)
    
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_create_syllabus_document_is_valid_docx(document_generator, sample_materials):
    """Test that the generated document is a valid DOCX file"""
    result = document_generator.create_syllabus_document(sample_materials)
    
    # Try to open the document - should not raise an exception
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    # Verify document has content
    assert len(doc.paragraphs) > 0


def test_document_contains_title(document_generator, sample_materials):
    """Test that the document contains the title"""
    result = document_generator.create_syllabus_document(sample_materials)
    
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    # Check for title in the document
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Silabus Kursus' in text_content


def test_document_contains_course_type(document_generator, sample_materials):
    """Test that the document contains the course type"""
    result = document_generator.create_syllabus_document(sample_materials)
    
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Jenis Kursus: B2B' in text_content


def test_document_contains_organization_section(document_generator, sample_materials):
    """Test that the document contains organization profile section"""
    result = document_generator.create_syllabus_document(sample_materials)
    
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Profil Organisasi' in text_content
    assert 'PT Example adalah perusahaan teknologi terkemuka' in text_content
    assert 'Perusahaan fokus pada solusi enterprise' in text_content


def test_document_contains_tlo_section(document_generator, sample_materials):
    """Test that the document contains TLO section"""
    result = document_generator.create_syllabus_document(sample_materials)
    
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Terminal Learning Objectives (TLO)' in text_content
    assert 'Peserta mampu memahami konsep dasar bisnis B2B' in text_content
    assert 'Peserta mampu mengimplementasikan strategi penjualan B2B' in text_content


def test_document_contains_performance_section(document_generator, sample_materials):
    """Test that the document contains performance section"""
    result = document_generator.create_syllabus_document(sample_materials)
    
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Performance Objectives' in text_content
    assert 'Mengidentifikasi karakteristik pasar B2B' in text_content


def test_document_contains_elo_section(document_generator, sample_materials):
    """Test that the document contains ELO section"""
    result = document_generator.create_syllabus_document(sample_materials)
    
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Enabling Learning Objectives (ELO)' in text_content
    assert 'Menjelaskan perbedaan antara pasar B2B dan B2C' in text_content


def test_document_with_empty_tlos(document_generator):
    """Test document generation with no TLOs selected"""
    materials = SyllabusMaterials(
        organization_summary="Test summary",
        organization_context="Test context",
        selected_tlos=[],
        selected_performances=[],
        selected_elos=[],
        course_type="Tech"
    )
    
    result = document_generator.create_syllabus_document(materials)
    
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Tidak ada TLO yang dipilih' in text_content


def test_document_with_empty_performances(document_generator):
    """Test document generation with no performances selected"""
    materials = SyllabusMaterials(
        organization_summary="Test summary",
        organization_context="Test context",
        selected_tlos=[],
        selected_performances=[],
        selected_elos=[],
        course_type="Tech"
    )
    
    result = document_generator.create_syllabus_document(materials)
    
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Tidak ada performance objectives yang dipilih' in text_content


def test_document_with_empty_elos(document_generator):
    """Test document generation with no ELOs selected"""
    materials = SyllabusMaterials(
        organization_summary="Test summary",
        organization_context="Test context",
        selected_tlos=[],
        selected_performances=[],
        selected_elos=[],
        course_type="Tech"
    )
    
    result = document_generator.create_syllabus_document(materials)
    
    doc_bytes = BytesIO(result)
    doc = Document(doc_bytes)
    
    text_content = '\n'.join([p.text for p in doc.paragraphs])
    assert 'Tidak ada ELO yang dipilih' in text_content
