"""
Export utilities: DataFrame → CSV, XLSX, DOCX, PPTX, PDF (in-memory bytes).
"""
import io
import os
import pandas as pd

# Path ke logo Telkom (relatif terhadap file ini)
_LOGO_PATH = os.path.join(os.path.dirname(__file__), "telkom_logo.png")


# ── Watermark helpers ─────────────────────────────────────────────────────────

def _pdf_watermark(canvas_obj, doc_obj):
    """ReportLab canvas callback — gambar logo Telkom diagonal semi-transparan di tengah halaman."""
    from reportlab.lib.units import cm
    try:
        canvas_obj.saveState()
        canvas_obj.setFillAlpha(0.07)
        w, h = doc_obj.pagesize
        logo_w = 10 * cm
        logo_h = 10 * cm
        canvas_obj.drawImage(
            _LOGO_PATH,
            (w - logo_w) / 2, (h - logo_h) / 2,
            width=logo_w, height=logo_h,
            preserveAspectRatio=True, mask="auto",
        )
        canvas_obj.restoreState()
    except Exception:
        pass  # Jika logo tidak ditemukan, lanjut tanpa watermark


def _docx_add_logo_header(doc):
    """Tambahkan logo Telkom kecil di header setiap section DOCX."""
    from docx.shared import Inches
    try:
        section = doc.sections[0]
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        run = p.add_run()
        run.add_picture(_LOGO_PATH, width=Inches(0.9))
    except Exception:
        pass  # Jika logo tidak ditemukan, lanjut tanpa header logo


def to_csv(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def to_xlsx(df: pd.DataFrame) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Data")
    return buf.getvalue()


def to_docx(df: pd.DataFrame, title: str = "Export") -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    _docx_add_logo_header(doc)
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val) if val is not None else ""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(df: pd.DataFrame, title: str = "Export") -> bytes:
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), leftMargin=1*cm, rightMargin=1*cm,
                            topMargin=1.5*cm, bottomMargin=1.5*cm)
    styles = getSampleStyleSheet()
    elements = [Paragraph(title, styles["Title"]), Spacer(1, 0.4*cm)]

    cols = list(df.columns)
    data = [cols] + [[str(v) if v is not None else "" for v in row] for _, row in df.iterrows()]

    col_width = (landscape(A4)[0] - 2*cm) / len(cols)
    tbl = Table(data, colWidths=[col_width] * len(cols), repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a3c6e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4fa")]),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("WORDWRAP", (0, 0), (-1, -1), True),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    elements.append(tbl)
    doc.build(elements, onFirstPage=_pdf_watermark, onLaterPages=_pdf_watermark)
    return buf.getvalue()


# ── Syllabus plain-text export ────────────────────────────────────────────────

def syllabus_to_text(final: dict) -> str:
    """Convert finalized syllabus dict to structured plain text."""
    profile = final.get("org_profile", {})
    course_type = final.get("course_type", "")
    lines = []

    lines.append("=" * 60)
    lines.append("DOKUMEN SILABUS PELATIHAN")
    lines.append("=" * 60)
    lines.append("")

    # Org profile
    lines.append("PROFIL ORGANISASI")
    lines.append("-" * 40)
    lines.append(f"Perusahaan  : {profile.get('organization_name', '-')}")
    lines.append(f"Industri    : {profile.get('industry', '-')}")
    lines.append(f"Visi        : {profile.get('vision', '-')}")
    lines.append(f"Misi        : {profile.get('mission', '-')}")
    lines.append("")
    lines.append("Prioritas Strategis:")
    for p in profile.get("strategic_priorities", []):
        lines.append(f"  - {p}")
    lines.append("")
    lines.append("Kompetensi Inti:")
    for c in profile.get("core_competencies", []):
        lines.append(f"  - {c}")
    lines.append("")
    lines.append(f"Konteks Pembelajaran:")
    lines.append(f"  {profile.get('learning_context', '-')}")
    lines.append("")

    lines.append(f"Tipe Course : {course_type}")
    if final.get("levels_covered"):
        lines.append(f"Level      : {' → '.join(final['levels_covered'])}")
    if final.get("current_condition"):
        lines.append(f"Kondisi Saat Ini    : {final['current_condition']}")
    if final.get("desired_condition"):
        lines.append(f"Kondisi Diinginkan  : {final['desired_condition']}")
    lines.append("")

    # TLOs
    lines.append("=" * 60)
    lines.append("TERMINAL LEARNING OBJECTIVES (TLO)")
    lines.append("=" * 60)
    for t in final.get("tlos", []):
        lines.append(f"\nTLO {t.get('tlo_number', '')}.")
        lines.append(f"  {t.get('tlo', '')}")
        lines.append(f"  Rationale: {t.get('rationale', '')}")

    lines.append("")

    # Performance Objectives
    lines.append("=" * 60)
    lines.append("PERFORMANCE OBJECTIVES")
    lines.append("=" * 60)
    for p in final.get("performance_objectives", []):
        lines.append(f"\nPO {p.get('perf_number', '')}. [{p.get('related_tlo', '')}]")
        lines.append(f"  {p.get('performance_objective', '')}")
        lines.append(f"  Kondisi  : {p.get('condition', '')}")
        lines.append(f"  Standar  : {p.get('standard', '')}")

    lines.append("")

    # ELOs
    lines.append("=" * 60)
    lines.append("ENABLING LEARNING OBJECTIVES (ELO)")
    lines.append("=" * 60)
    total_dur = sum(e.get("duration_minutes", 0) for e in final.get("elos", []))
    for e in final.get("elos", []):
        lines.append(f"\nELO {e.get('elo_number', '')}. [{e.get('related_performance', '')}]")
        lines.append(f"  {e.get('elo', '')}")
        lines.append(f"  Bloom Level     : {e.get('bloom_level', '')}")
        lines.append(f"  Metode Delivery : {e.get('delivery_method', '')}")
        lines.append(f"  Durasi          : {e.get('duration_minutes', '')} menit")

    lines.append("")
    lines.append("-" * 60)
    lines.append(f"Total Estimasi Durasi: {total_dur} menit")
    lines.append("=" * 60)

    return "\n".join(lines)


def syllabus_to_docx(final: dict, title: str = "Silabus Pelatihan") -> bytes:
    """Export syllabus as plain-text DOCX (no tables)."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    profile = final.get("org_profile", {})
    course_type = final.get("course_type", "")

    doc = Document()
    _docx_add_logo_header(doc)
    doc.add_heading(title, level=0)

    # Org profile section
    doc.add_heading("Profil Perusahaan", level=1)
    fields = [
        ("Perusahaan", profile.get("organization_name", "-")),
        ("Industri", profile.get("industry", "-")),
        ("Visi", profile.get("vision", "-")),
        ("Misi", profile.get("mission", "-")),
    ]
    for label, val in fields:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(val)

    doc.add_paragraph("Prioritas Strategis:").runs[0].bold = True
    for item in profile.get("strategic_priorities", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Kompetensi Inti:").runs[0].bold = True
    for item in profile.get("core_competencies", []):
        doc.add_paragraph(item, style="List Bullet")

    p = doc.add_paragraph()
    p.add_run("Konteks Pembelajaran: ").bold = True
    p.add_run(profile.get("learning_context", "-"))

    p = doc.add_paragraph()
    p.add_run("Tipe Course: ").bold = True
    p.add_run(course_type)

    if final.get("levels_covered"):
        p = doc.add_paragraph()
        p.add_run("Level yang Dicakup: ").bold = True
        p.add_run(" → ".join(final["levels_covered"]))

    if final.get("current_condition"):
        p = doc.add_paragraph()
        p.add_run("Kondisi Saat Ini: ").bold = True
        p.add_run(final["current_condition"])
    if final.get("desired_condition"):
        p = doc.add_paragraph()
        p.add_run("Kondisi yang Diinginkan: ").bold = True
        p.add_run(final["desired_condition"])

    # TLOs
    doc.add_heading("Terminal Learning Objectives (TLO)", level=1)
    for t in final.get("tlos", []):
        doc.add_heading(f"TLO {t.get('tlo_number', '')}.", level=2)
        doc.add_paragraph(t.get("tlo", ""))
        p = doc.add_paragraph()
        p.add_run("Rationale: ").bold = True
        p.add_run(t.get("rationale", ""))

    # Performance Objectives
    doc.add_heading("Performance Objectives", level=1)
    for po in final.get("performance_objectives", []):
        doc.add_heading(f"PO {po.get('perf_number', '')}. [{po.get('related_tlo', '')}]", level=2)
        doc.add_paragraph(po.get("performance_objective", ""))
        p = doc.add_paragraph()
        p.add_run("Kondisi: ").bold = True
        p.add_run(po.get("condition", ""))
        p = doc.add_paragraph()
        p.add_run("Standar: ").bold = True
        p.add_run(po.get("standard", ""))

    # ELOs
    doc.add_heading("Enabling Learning Objectives (ELO)", level=1)
    total_dur = sum(e.get("duration_minutes", 0) for e in final.get("elos", []))
    for e in final.get("elos", []):
        doc.add_heading(f"ELO {e.get('elo_number', '')}. [{e.get('related_performance', '')}]", level=2)
        doc.add_paragraph(e.get("elo", ""))
        p = doc.add_paragraph()
        p.add_run("Bloom Level: ").bold = True
        p.add_run(e.get("bloom_level", ""))
        p.add_run("   |   Metode: ").bold = True
        p.add_run(e.get("delivery_method", ""))
        p.add_run("   |   Durasi: ").bold = True
        p.add_run(f"{e.get('duration_minutes', '')} menit")

    doc.add_paragraph(f"\nTotal Estimasi Durasi: {total_dur} menit")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def syllabus_to_pdf(final: dict, title: str = "Silabus Pelatihan") -> bytes:
    """Export syllabus as plain-text PDF (no tables)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    normal = styles["Normal"]
    bold_style = ParagraphStyle("bold", parent=normal, fontName="Helvetica-Bold")
    label_style = ParagraphStyle("label", parent=normal, fontName="Helvetica-Bold",
                                 textColor=colors.HexColor("#1a3c6e"))

    profile = final.get("org_profile", {})
    course_type = final.get("course_type", "")
    elements = []

    def add(text, style=normal, space=0.2):
        elements.append(Paragraph(text, style))
        elements.append(Spacer(1, space * cm))

    def hr():
        elements.append(HRFlowable(width="100%", thickness=0.5,
                                   color=colors.HexColor("#cccccc")))
        elements.append(Spacer(1, 0.2*cm))

    add(title, styles["Title"], 0.4)
    hr()

    add("Profil Perusahaan", h1, 0.1)
    add(f"<b>Perusahaan:</b> {profile.get('organization_name', '-')}")
    add(f"<b>Industri:</b> {profile.get('industry', '-')}")
    add(f"<b>Visi:</b> {profile.get('vision', '-')}")
    add(f"<b>Misi:</b> {profile.get('mission', '-')}")
    add("<b>Prioritas Strategis:</b>")
    for p in profile.get("strategic_priorities", []):
        add(f"• {p}")
    add("<b>Kompetensi Inti:</b>")
    for c in profile.get("core_competencies", []):
        add(f"• {c}")
    add(f"<b>Konteks Pembelajaran:</b> {profile.get('learning_context', '-')}")
    add(f"<b>Tipe Course:</b> {course_type}")
    if final.get("levels_covered"):
        add(f"<b>Level yang Dicakup:</b> {' → '.join(final['levels_covered'])}")
    if final.get("current_condition"):
        add(f"<b>Kondisi Saat Ini:</b> {final['current_condition']}")
    if final.get("desired_condition"):
        add(f"<b>Kondisi yang Diinginkan:</b> {final['desired_condition']}")
    hr()

    add("Terminal Learning Objectives (TLO)", h1, 0.1)
    for t in final.get("tlos", []):
        add(f"<b>TLO {t.get('tlo_number', '')}.</b>", bold_style, 0.05)
        add(t.get("tlo", ""))
        add(f"<i>Rationale: {t.get('rationale', '')}</i>", space=0.3)
    hr()

    add("Performance Objectives", h1, 0.1)
    for po in final.get("performance_objectives", []):
        add(f"<b>PO {po.get('perf_number', '')}.</b> [{po.get('related_tlo', '')}]", bold_style, 0.05)
        add(po.get("performance_objective", ""))
        add(f"<b>Kondisi:</b> {po.get('condition', '')}")
        add(f"<b>Standar:</b> {po.get('standard', '')}", space=0.3)
    hr()

    add("Enabling Learning Objectives (ELO)", h1, 0.1)
    total_dur = sum(e.get("duration_minutes", 0) for e in final.get("elos", []))
    for e in final.get("elos", []):
        add(f"<b>ELO {e.get('elo_number', '')}.</b> [{e.get('related_performance', '')}]", bold_style, 0.05)
        add(e.get("elo", ""))
        add(f"<b>Bloom:</b> {e.get('bloom_level', '')}  |  "
            f"<b>Metode:</b> {e.get('delivery_method', '')}  |  "
            f"<b>Durasi:</b> {e.get('duration_minutes', '')} menit", space=0.3)
    hr()
    add(f"<b>Total Estimasi Durasi: {total_dur} menit</b>")

    doc.build(elements, onFirstPage=_pdf_watermark, onLaterPages=_pdf_watermark)
    return buf.getvalue()


# ── Decompose plain-text export ───────────────────────────────────────────────

def decompose_to_text(modules: list, source_name: str = "") -> str:
    lines = []
    lines.append("=" * 60)
    lines.append("DEKOMPOSISI MODUL MIKRO")
    if source_name:
        lines.append(f"Sumber: {source_name}")
    total = sum(m.get("duration_minutes", 0) for m in modules)
    lines.append(f"Total Modul: {len(modules)}  |  Total Durasi: {total} menit")
    lines.append("=" * 60)

    for m in modules:
        lines.append("")
        lines.append(f"Modul {m.get('module_number', '')}. {m.get('title', '')}")
        lines.append("-" * 40)
        lines.append(f"Tujuan    : {m.get('specific_objective', '')}")
        lines.append(f"Ringkasan : {m.get('content_summary', '')}")
        lines.append(f"Format    : {m.get('delivery_format', '')}")
        lines.append(f"Durasi    : {m.get('duration_minutes', '')} menit")

    lines.append("")
    lines.append("=" * 60)
    return "\n".join(lines)


def decompose_to_docx(modules: list, source_name: str = "") -> bytes:
    from docx import Document
    doc = Document()
    _docx_add_logo_header(doc)
    title = f"Dekomposisi Modul Mikro"
    if source_name:
        title += f" — {source_name}"
    doc.add_heading(title, level=0)

    total = sum(m.get("duration_minutes", 0) for m in modules)
    p = doc.add_paragraph()
    p.add_run(f"Total Modul: {len(modules)}  |  Total Durasi: {total} menit")

    for m in modules:
        doc.add_heading(f"Modul {m.get('module_number', '')}. {m.get('title', '')}", level=1)
        fields = [
            ("Tujuan", m.get("specific_objective", "")),
            ("Ringkasan", m.get("content_summary", "")),
            ("Format", m.get("delivery_format", "")),
            ("Durasi", f"{m.get('duration_minutes', '')} menit"),
        ]
        for label, val in fields:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(val)

    doc.add_page_break()
    doc.add_heading("Timeline Penyelesaian Modul", level=1)
    timeline_to_docx_section(doc, modules)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def decompose_to_pdf(modules: list, source_name: str = "") -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    bold_style = ParagraphStyle("bold", parent=styles["Normal"], fontName="Helvetica-Bold")
    elements = []

    title = "Dekomposisi Modul Mikro"
    if source_name:
        title += f" — {source_name}"
    elements.append(Paragraph(title, styles["Title"]))
    elements.append(Spacer(1, 0.3*cm))

    total = sum(m.get("duration_minutes", 0) for m in modules)
    elements.append(Paragraph(f"Total Modul: {len(modules)}  |  Total Durasi: {total} menit", styles["Normal"]))
    elements.append(Spacer(1, 0.4*cm))

    for m in modules:
        elements.append(Paragraph(f"Modul {m.get('module_number', '')}. {m.get('title', '')}", styles["Heading1"]))
        elements.append(Paragraph(f"<b>Tujuan:</b> {m.get('specific_objective', '')}", styles["Normal"]))
        elements.append(Spacer(1, 0.15*cm))
        elements.append(Paragraph(f"<b>Ringkasan:</b> {m.get('content_summary', '')}", styles["Normal"]))
        elements.append(Spacer(1, 0.15*cm))
        elements.append(Paragraph(
            f"<b>Format:</b> {m.get('delivery_format', '')}  |  "
            f"<b>Durasi:</b> {m.get('duration_minutes', '')} menit", styles["Normal"]))
        elements.append(Spacer(1, 0.3*cm))
        elements.append(HRFlowable(width="100%", thickness=0.4, color=colors.HexColor("#cccccc")))
        elements.append(Spacer(1, 0.2*cm))

    from reportlab.platypus import PageBreak
    elements.append(PageBreak())
    elements.append(Paragraph("Timeline Penyelesaian Modul", styles["Heading1"]))
    elements.append(Spacer(1, 0.3*cm))
    elements.extend(timeline_to_pdf_elements(modules, styles))

    doc.build(elements, onFirstPage=_pdf_watermark, onLaterPages=_pdf_watermark)
    return buf.getvalue()


# ── Recommend plain-text export ───────────────────────────────────────────────

def recommend_to_text(recs: list, participant: str = "", gap: str = "") -> str:
    lines = []
    lines.append("=" * 60)
    lines.append("REKOMENDASI LEARNING PATH PERSONAL")
    if participant:
        lines.append(f"Peserta : {participant}")
    if gap:
        lines.append(f"Gap     : {gap}")
    total = sum(r.get("estimated_duration_minutes", 0) for r in recs)
    lines.append(f"Total Estimasi Durasi: {total} menit")
    lines.append("=" * 60)

    for r in recs:
        priority = r.get("priority", "")
        lines.append("")
        lines.append(f"#{r.get('rank', '')}. {r.get('module_title', '')}  [{priority}]")
        lines.append("-" * 40)
        lines.append(f"Alasan Relevansi : {r.get('relevance_reason', '')}")
        lines.append(f"Estimasi Durasi  : {r.get('estimated_duration_minutes', '')} menit")

    lines.append("")
    lines.append("=" * 60)
    return "\n".join(lines)


def recommend_to_docx(recs: list, participant: str = "", gap: str = "") -> bytes:
    from docx import Document
    doc = Document()
    _docx_add_logo_header(doc)
    doc.add_heading("Rekomendasi Learning Path Personal", level=0)

    if participant:
        p = doc.add_paragraph()
        p.add_run("Peserta: ").bold = True
        p.add_run(participant)
    if gap:
        p = doc.add_paragraph()
        p.add_run("Gap Kompetensi: ").bold = True
        p.add_run(gap)

    total = sum(r.get("estimated_duration_minutes", 0) for r in recs)
    doc.add_paragraph(f"Total Estimasi Durasi: {total} menit")

    for r in recs:
        doc.add_heading(f"#{r.get('rank', '')}. {r.get('module_title', '')}  [{r.get('priority', '')}]", level=1)
        p = doc.add_paragraph()
        p.add_run("Alasan Relevansi: ").bold = True
        p.add_run(r.get("relevance_reason", ""))
        p = doc.add_paragraph()
        p.add_run("Estimasi Durasi: ").bold = True
        p.add_run(f"{r.get('estimated_duration_minutes', '')} menit")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def recommend_to_pdf(recs: list, participant: str = "", gap: str = "") -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Rekomendasi Learning Path Personal", styles["Title"]))
    elements.append(Spacer(1, 0.3*cm))
    if participant:
        elements.append(Paragraph(f"<b>Peserta:</b> {participant}", styles["Normal"]))
    if gap:
        elements.append(Paragraph(f"<b>Gap Kompetensi:</b> {gap}", styles["Normal"]))
    total = sum(r.get("estimated_duration_minutes", 0) for r in recs)
    elements.append(Paragraph(f"<b>Total Estimasi Durasi:</b> {total} menit", styles["Normal"]))
    elements.append(Spacer(1, 0.4*cm))

    priority_colors = {"High": "#c0392b", "Medium": "#e67e22", "Low": "#27ae60"}
    for r in recs:
        color = priority_colors.get(r.get("priority", ""), "#333333")
        elements.append(Paragraph(
            f'<font color="{color}"><b>#{r.get("rank", "")}. {r.get("module_title", "")}  [{r.get("priority", "")}]</b></font>',
            styles["Heading1"]))
        elements.append(Paragraph(f"<b>Alasan Relevansi:</b> {r.get('relevance_reason', '')}", styles["Normal"]))
        elements.append(Spacer(1, 0.15*cm))
        elements.append(Paragraph(f"<b>Estimasi Durasi:</b> {r.get('estimated_duration_minutes', '')} menit", styles["Normal"]))
        elements.append(Spacer(1, 0.3*cm))
        elements.append(HRFlowable(width="100%", thickness=0.4, color=colors.HexColor("#cccccc")))
        elements.append(Spacer(1, 0.2*cm))

    doc.build(elements, onFirstPage=_pdf_watermark, onLaterPages=_pdf_watermark)
    return buf.getvalue()


# ── Combined syllabus + decompose export ─────────────────────────────────────

def combined_to_docx(syllabus_data: dict, modules: list) -> bytes:
    """Export silabus + modul mikro dalam satu DOCX."""
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    _docx_add_logo_header(doc)

    output = syllabus_data.get("output_json", {})
    profile = output.get("org_profile", {})
    course_type = output.get("course_type", syllabus_data.get("topic", ""))

    doc.add_heading(f"Silabus & Modul Mikro: {course_type}", level=0)

    # ── Bagian 1: Silabus ──
    doc.add_heading("Bagian 1 — Silabus Pelatihan", level=1)

    doc.add_heading("Profil Perusahaan", level=2)
    for label, val in [
        ("Perusahaan", profile.get("organization_name", "-")),
        ("Industri", profile.get("industry", "-")),
        ("Visi", profile.get("vision", "-")),
        ("Misi", profile.get("mission", "-")),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(val)

    p = doc.add_paragraph()
    p.add_run("Tipe Course: ").bold = True
    p.add_run(course_type)
    if output.get("levels_covered"):
        p = doc.add_paragraph()
        p.add_run("Level: ").bold = True
        p.add_run(" → ".join(output["levels_covered"]))
    if output.get("current_condition"):
        p = doc.add_paragraph()
        p.add_run("Kondisi Saat Ini: ").bold = True
        p.add_run(output["current_condition"])
    if output.get("desired_condition"):
        p = doc.add_paragraph()
        p.add_run("Kondisi yang Diinginkan: ").bold = True
        p.add_run(output["desired_condition"])

    doc.add_heading("Terminal Learning Objectives (TLO)", level=2)
    for t in output.get("tlos", []):
        doc.add_heading(f"TLO {t.get('tlo_number', '')}.", level=3)
        doc.add_paragraph(t.get("tlo", ""))
        p = doc.add_paragraph()
        p.add_run("Rationale: ").bold = True
        p.add_run(t.get("rationale", ""))

    doc.add_heading("Performance Objectives", level=2)
    for po in output.get("performance_objectives", []):
        doc.add_heading(f"PO {po.get('perf_number', '')}. [{po.get('related_tlo', '')}]", level=3)
        doc.add_paragraph(po.get("performance_objective", ""))
        p = doc.add_paragraph()
        p.add_run("Kondisi: ").bold = True
        p.add_run(po.get("condition", ""))
        p = doc.add_paragraph()
        p.add_run("Standar: ").bold = True
        p.add_run(po.get("standard", ""))

    doc.add_heading("Enabling Learning Objectives (ELO)", level=2)
    for e in output.get("elos", []):
        doc.add_heading(f"ELO {e.get('elo_number', '')}. [{e.get('related_performance', '')}]", level=3)
        doc.add_paragraph(e.get("elo", ""))
        p = doc.add_paragraph()
        p.add_run("Bloom: ").bold = True
        p.add_run(e.get("bloom_level", ""))
        p.add_run("  |  Metode: ").bold = True
        p.add_run(e.get("delivery_method", ""))
        p.add_run("  |  Durasi: ").bold = True
        p.add_run(f"{e.get('duration_minutes', '')} menit")

    # ── Bagian 2: Modul Mikro ──
    doc.add_page_break()
    doc.add_heading("Bagian 2 — Dekomposisi Modul Mikro", level=1)
    total = sum(m.get("duration_minutes", 0) for m in modules)
    p = doc.add_paragraph()
    p.add_run(f"Total Modul: {len(modules)}  |  Total Durasi: {total} menit")

    for m in modules:
        doc.add_heading(f"Modul {m.get('module_number', '')}. {m.get('title', '')}", level=2)
        fields = [
            ("Tujuan", m.get("specific_objective", "")),
            ("Ringkasan", m.get("content_summary", "")),
            ("Format", m.get("delivery_format", "")),
            ("Durasi", f"{m.get('duration_minutes', '')} menit"),
        ]
        if m.get("related_elo"):
            fields.append(("Berdasarkan", m.get("related_elo", "")))
        for label, val in fields:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(val)

    # ── Timeline ──
    doc.add_page_break()
    doc.add_heading("Bagian 3 — Timeline Penyelesaian Modul", level=1)
    timeline_to_docx_section(doc, modules)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def combined_to_pdf(syllabus_data: dict, modules: list) -> bytes:
    """Export silabus + modul mikro dalam satu PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm

    output = syllabus_data.get("output_json", {})
    profile = output.get("org_profile", {})
    course_type = output.get("course_type", syllabus_data.get("topic", ""))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    bold_style = ParagraphStyle("bold", parent=normal, fontName="Helvetica-Bold")
    elements = []

    def add(text, style=normal, space=0.2):
        elements.append(Paragraph(text, style))
        elements.append(Spacer(1, space * cm))

    def hr():
        elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))
        elements.append(Spacer(1, 0.2*cm))

    # Cover
    add(f"Silabus & Modul Mikro: {course_type}", styles["Title"], 0.4)
    hr()

    # Bagian 1
    add("Bagian 1 — Silabus Pelatihan", styles["Heading1"], 0.2)
    add(f"<b>Perusahaan:</b> {profile.get('organization_name', '-')}")
    add(f"<b>Industri:</b> {profile.get('industry', '-')}")
    add(f"<b>Visi:</b> {profile.get('vision', '-')}")
    add(f"<b>Misi:</b> {profile.get('mission', '-')}")
    add(f"<b>Tipe Course:</b> {course_type}")
    if output.get("levels_covered"):
        add(f"<b>Level:</b> {' → '.join(output['levels_covered'])}")
    if output.get("current_condition"):
        add(f"<b>Kondisi Saat Ini:</b> {output['current_condition']}")
    if output.get("desired_condition"):
        add(f"<b>Kondisi yang Diinginkan:</b> {output['desired_condition']}")
    hr()

    add("Terminal Learning Objectives (TLO)", styles["Heading2"], 0.1)
    for t in output.get("tlos", []):
        add(f"<b>TLO {t.get('tlo_number', '')}.</b> {t.get('tlo', '')}", bold_style, 0.05)
        add(f"<i>Rationale: {t.get('rationale', '')}</i>", space=0.25)
    hr()

    add("Performance Objectives", styles["Heading2"], 0.1)
    for po in output.get("performance_objectives", []):
        add(f"<b>PO {po.get('perf_number', '')}.</b> [{po.get('related_tlo', '')}] {po.get('performance_objective', '')}", bold_style, 0.05)
        add(f"<b>Kondisi:</b> {po.get('condition', '')}  |  <b>Standar:</b> {po.get('standard', '')}", space=0.25)
    hr()

    add("Enabling Learning Objectives (ELO)", styles["Heading2"], 0.1)
    for e in output.get("elos", []):
        add(f"<b>ELO {e.get('elo_number', '')}.</b> [{e.get('related_performance', '')}] {e.get('elo', '')}", bold_style, 0.05)
        add(f"<b>Bloom:</b> {e.get('bloom_level', '')}  |  <b>Metode:</b> {e.get('delivery_method', '')}  |  <b>Durasi:</b> {e.get('duration_minutes', '')} menit", space=0.25)
    hr()

    # Bagian 2
    elements.append(PageBreak())
    add("Bagian 2 — Dekomposisi Modul Mikro", styles["Heading1"], 0.2)
    total = sum(m.get("duration_minutes", 0) for m in modules)
    add(f"<b>Total Modul:</b> {len(modules)}  |  <b>Total Durasi:</b> {total} menit")
    hr()

    for m in modules:
        add(f"<b>Modul {m.get('module_number', '')}. {m.get('title', '')}</b>", styles["Heading2"], 0.05)
        add(f"<b>Tujuan:</b> {m.get('specific_objective', '')}")
        add(f"<b>Ringkasan:</b> {m.get('content_summary', '')}")
        add(f"<b>Format:</b> {m.get('delivery_format', '')}  |  <b>Durasi:</b> {m.get('duration_minutes', '')} menit")
        if m.get("related_elo"):
            add(f"<b>Berdasarkan:</b> {m.get('related_elo', '')}")
        elements.append(Spacer(1, 0.3*cm))
        hr()

    # Bagian 3: Timeline
    elements.append(PageBreak())
    elements.append(Paragraph("Bagian 3 — Timeline Penyelesaian Modul", styles["Heading1"]))
    elements.append(Spacer(1, 0.3*cm))
    elements.extend(timeline_to_pdf_elements(modules, styles))

    doc.build(elements, onFirstPage=_pdf_watermark, onLaterPages=_pdf_watermark)
    return buf.getvalue()


# ── Timeline generator ────────────────────────────────────────────────────────

def build_timeline(modules: list) -> dict:
    """
    Distribute modules into two timeline versions:
    - Short: 2 modules per week
    - Long:  1 module per week
    Returns dict with 'short' and 'long' as list of week dicts.
    """
    def distribute(modules, per_week):
        weeks = []
        for i in range(0, len(modules), per_week):
            chunk = modules[i:i + per_week]
            weeks.append({
                "week": len(weeks) + 1,
                "modules": [f"Modul {m.get('module_number', i+1)}: {m.get('title', '')}" for m in chunk],
                "module_count": len(chunk),
                "total_duration_minutes": sum(m.get("duration_minutes", 0) for m in chunk),
            })
        return weeks

    return {
        "short": distribute(modules, 2),  # 2 modul/minggu
        "long":  distribute(modules, 1),  # 1 modul/minggu
    }


def timeline_to_df(weeks: list) -> "pd.DataFrame":
    rows = []
    for w in weeks:
        rows.append({
            "Minggu": f"Minggu {w['week']}",
            "Jumlah Modul": w["module_count"],
            "Total Durasi (menit)": w["total_duration_minutes"],
            "Modul yang Dipelajari": "\n".join(w["modules"]),
        })
    return pd.DataFrame(rows)


def timeline_to_docx_section(doc, modules: list):
    """Append timeline tables (short + long) to an existing DOCX document."""
    from docx.shared import Pt
    timeline = build_timeline(modules)

    for label, weeks in [("Versi Singkat (2 Modul per Minggu)", timeline["short"]),
                          ("Versi Lama (1 Modul per Minggu)", timeline["long"])]:
        doc.add_heading(label, level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Minggu", "Jumlah Modul", "Total Durasi (menit)", "Modul yang Dipelajari"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        for w in weeks:
            row = table.add_row().cells
            row[0].text = f"Minggu {w['week']}"
            row[1].text = str(w["module_count"])
            row[2].text = str(w["total_duration_minutes"])
            row[3].text = "\n".join(w["modules"])
        doc.add_paragraph()


def timeline_to_pdf_elements(modules: list, styles):
    """Return ReportLab flowables for timeline tables (short + long)."""
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    timeline = build_timeline(modules)
    elements = []

    col_widths = [2.5*cm, 2.5*cm, 3.5*cm, 10.5*cm]

    for label, weeks in [("Versi Singkat (2 Modul per Minggu)", timeline["short"]),
                          ("Versi Lama (1 Modul per Minggu)", timeline["long"])]:
        elements.append(Paragraph(label, styles["Heading2"]))
        elements.append(Spacer(1, 0.2*cm))

        data = [["Minggu", "Jml Modul", "Durasi (mnt)", "Modul yang Dipelajari"]]
        for w in weeks:
            data.append([
                f"Minggu {w['week']}",
                str(w["module_count"]),
                str(w["total_duration_minutes"]),
                "\n".join(w["modules"]),
            ])

        tbl = Table(data, colWidths=col_widths, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a3c6e")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4fa")]),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        elements.append(tbl)
        elements.append(Spacer(1, 0.5*cm))

    return elements


# ── Decompose manual (without syllabus) export — includes org profile ─────────

def decompose_manual_to_docx(modules: list, manual_meta: dict) -> bytes:
    """Export modul mikro + profil perusahaan (mode tanpa silabus) sebagai DOCX."""
    from docx import Document
    doc = Document()
    _docx_add_logo_header(doc)

    profile = manual_meta.get("org_profile", {})
    course_type = manual_meta.get("course_type", "")
    levels = manual_meta.get("levels_covered", [])

    doc.add_heading(f"Modul Mikro: {course_type}", level=0)

    # Profil perusahaan
    doc.add_heading("Profil Perusahaan & Course", level=1)
    for label, val in [
        ("Perusahaan", profile.get("organization_name", "-")),
        ("Industri", profile.get("industry", "-")),
        ("Visi", profile.get("vision", "-")),
        ("Misi", profile.get("mission", "-")),
        ("Tipe Course", course_type),
        ("Level", " → ".join(levels) if levels else "-"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(val)
    if profile.get("strategic_priorities"):
        doc.add_paragraph("Prioritas Strategis:").runs[0].bold = True
        for item in profile["strategic_priorities"]:
            doc.add_paragraph(item, style="List Bullet")
    if profile.get("learning_context"):
        p = doc.add_paragraph()
        p.add_run("Konteks Pembelajaran: ").bold = True
        p.add_run(profile["learning_context"])

    # Modul mikro
    doc.add_page_break()
    doc.add_heading("Daftar Modul Mikro", level=1)
    total = sum(m.get("duration_minutes", 0) for m in modules)
    p = doc.add_paragraph()
    p.add_run(f"Total Modul: {len(modules)}  |  Total Durasi: {total} menit")

    for m in modules:
        doc.add_heading(f"Modul {m.get('module_number', '')}. {m.get('title', '')}", level=2)
        for label, val in [
            ("Tujuan", m.get("specific_objective", "")),
            ("Ringkasan", m.get("content_summary", "")),
            ("Format", m.get("delivery_format", "")),
            ("Durasi", f"{m.get('duration_minutes', '')} menit"),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(val)
        if m.get("related_elo"):
            p = doc.add_paragraph()
            p.add_run("Berdasarkan: ").bold = True
            p.add_run(m["related_elo"])

    # Timeline
    doc.add_page_break()
    doc.add_heading("Timeline Penyelesaian Modul", level=1)
    timeline_to_docx_section(doc, modules)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def decompose_manual_to_pdf(modules: list, manual_meta: dict) -> bytes:
    """Export modul mikro + profil perusahaan (mode tanpa silabus) sebagai PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm

    profile = manual_meta.get("org_profile", {})
    course_type = manual_meta.get("course_type", "")
    levels = manual_meta.get("levels_covered", [])

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    bold_style = ParagraphStyle("bold", parent=normal, fontName="Helvetica-Bold")
    elements = []

    def add(text, style=normal, space=0.2):
        elements.append(Paragraph(text, style))
        elements.append(Spacer(1, space * cm))

    def hr():
        elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc")))
        elements.append(Spacer(1, 0.2*cm))

    add(f"Modul Mikro: {course_type}", styles["Title"], 0.4)
    hr()

    add("Profil Perusahaan & Course", styles["Heading1"], 0.1)
    add(f"<b>Perusahaan:</b> {profile.get('organization_name', '-')}")
    add(f"<b>Industri:</b> {profile.get('industry', '-')}")
    add(f"<b>Visi:</b> {profile.get('vision', '-')}")
    add(f"<b>Misi:</b> {profile.get('mission', '-')}")
    add(f"<b>Tipe Course:</b> {course_type}")
    add(f"<b>Level:</b> {' → '.join(levels) if levels else '-'}")
    if profile.get("strategic_priorities"):
        add("<b>Prioritas Strategis:</b>")
        for p in profile["strategic_priorities"]:
            add(f"• {p}")
    if profile.get("learning_context"):
        add(f"<b>Konteks Pembelajaran:</b> {profile['learning_context']}")
    hr()

    elements.append(PageBreak())
    add("Daftar Modul Mikro", styles["Heading1"], 0.2)
    total = sum(m.get("duration_minutes", 0) for m in modules)
    add(f"<b>Total Modul:</b> {len(modules)}  |  <b>Total Durasi:</b> {total} menit")
    hr()

    for m in modules:
        add(f"<b>Modul {m.get('module_number', '')}. {m.get('title', '')}</b>", styles["Heading2"], 0.05)
        add(f"<b>Tujuan:</b> {m.get('specific_objective', '')}")
        add(f"<b>Ringkasan:</b> {m.get('content_summary', '')}")
        add(f"<b>Format:</b> {m.get('delivery_format', '')}  |  <b>Durasi:</b> {m.get('duration_minutes', '')} menit")
        if m.get("related_elo"):
            add(f"<b>Berdasarkan:</b> {m.get('related_elo', '')}")
        elements.append(Spacer(1, 0.3*cm))
        hr()

    elements.append(PageBreak())
    elements.append(Paragraph("Timeline Penyelesaian Modul", styles["Heading1"]))
    elements.append(Spacer(1, 0.3*cm))
    elements.extend(timeline_to_pdf_elements(modules, styles))

    doc.build(elements, onFirstPage=_pdf_watermark, onLaterPages=_pdf_watermark)
    return buf.getvalue()


def decompose_manual_to_text(modules: list, manual_meta: dict) -> str:
    """Plain text export: profil perusahaan + modul mikro (mode tanpa silabus)."""
    profile = manual_meta.get("org_profile", {})
    course_type = manual_meta.get("course_type", "")
    levels = manual_meta.get("levels_covered", [])
    lines = []

    lines.append("=" * 60)
    lines.append("MODUL MIKRO PELATIHAN")
    lines.append("=" * 60)
    lines.append("")
    lines.append("PROFIL PERUSAHAAN & COURSE")
    lines.append("-" * 40)
    lines.append(f"Perusahaan  : {profile.get('organization_name', '-')}")
    lines.append(f"Industri    : {profile.get('industry', '-')}")
    lines.append(f"Visi        : {profile.get('vision', '-')}")
    lines.append(f"Misi        : {profile.get('mission', '-')}")
    lines.append(f"Tipe Course : {course_type}")
    lines.append(f"Level       : {' → '.join(levels) if levels else '-'}")
    if profile.get("strategic_priorities"):
        lines.append("Prioritas Strategis:")
        for p in profile["strategic_priorities"]:
            lines.append(f"  - {p}")
    if profile.get("learning_context"):
        lines.append(f"Konteks Pembelajaran: {profile['learning_context']}")
    lines.append("")

    total = sum(m.get("duration_minutes", 0) for m in modules)
    lines.append("=" * 60)
    lines.append("DAFTAR MODUL MIKRO")
    lines.append(f"Total Modul: {len(modules)}  |  Total Durasi: {total} menit")
    lines.append("=" * 60)

    for m in modules:
        lines.append("")
        lines.append(f"Modul {m.get('module_number', '')}. {m.get('title', '')}")
        lines.append("-" * 40)
        lines.append(f"Tujuan    : {m.get('specific_objective', '')}")
        lines.append(f"Ringkasan : {m.get('content_summary', '')}")
        lines.append(f"Format    : {m.get('delivery_format', '')}")
        lines.append(f"Durasi    : {m.get('duration_minutes', '')} menit")
        if m.get("related_elo"):
            lines.append(f"Berdasarkan: {m.get('related_elo', '')}")

    lines.append("")
    lines.append("=" * 60)
    return "\n".join(lines)


def combined_to_text(syllabus_data: dict, modules: list) -> str:
    """Plain text export: silabus lengkap + modul mikro."""
    output = syllabus_data.get("output_json", {})
    # Reuse syllabus_to_text for the syllabus part
    syl_text = syllabus_to_text(output)
    # Then append modules
    total = sum(m.get("duration_minutes", 0) for m in modules)
    lines = [syl_text, "", "=" * 60, "DEKOMPOSISI MODUL MIKRO",
             f"Total Modul: {len(modules)}  |  Total Durasi: {total} menit",
             "=" * 60]
    for m in modules:
        lines.append("")
        lines.append(f"Modul {m.get('module_number', '')}. {m.get('title', '')}")
        lines.append("-" * 40)
        lines.append(f"Tujuan    : {m.get('specific_objective', '')}")
        lines.append(f"Ringkasan : {m.get('content_summary', '')}")
        lines.append(f"Format    : {m.get('delivery_format', '')}")
        lines.append(f"Durasi    : {m.get('duration_minutes', '')} menit")
        if m.get("related_elo"):
            lines.append(f"Berdasarkan: {m.get('related_elo', '')}")
    lines.append("")
    lines.append("=" * 60)
    return "\n".join(lines)
