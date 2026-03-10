"""Manual test script for DocumentGenerator"""

from src.processors.document_generator import DocumentGenerator
from src.models.entities import SyllabusMaterials, TLO, Performance, ELO
from docx import Document
from io import BytesIO


def test_document_generator():
    """Test the DocumentGenerator class"""
    print("Testing DocumentGenerator...")
    
    # Create sample data
    tlo1 = TLO(
        id="tlo-1",
        org_id="org-1",
        course_type="B2B",
        text="Peserta mampu memahami konsep dasar bisnis B2B",
        is_selected=True
    )
    
    tlo2 = TLO(
        id="tlo-2",
        org_id="org-1",
        course_type="B2B",
        text="Peserta mampu mengimplementasikan strategi penjualan B2B",
        is_selected=True
    )
    
    perf1 = Performance(
        id="perf-1",
        tlo_ids=["tlo-1"],
        text="Mengidentifikasi karakteristik pasar B2B",
        is_selected=True
    )
    
    perf2 = Performance(
        id="perf-2",
        tlo_ids=["tlo-2"],
        text="Merancang strategi pendekatan klien B2B",
        is_selected=True
    )
    
    elo1 = ELO(
        id="elo-1",
        performance_ids=["perf-1"],
        text="Menjelaskan perbedaan antara pasar B2B dan B2C",
        is_selected=True
    )
    
    elo2 = ELO(
        id="elo-2",
        performance_ids=["perf-1"],
        text="Menganalisis kebutuhan pelanggan B2B",
        is_selected=True
    )
    
    elo3 = ELO(
        id="elo-3",
        performance_ids=["perf-2"],
        text="Menyusun proposal nilai untuk klien B2B",
        is_selected=True
    )
    
    materials = SyllabusMaterials(
        organization_summary="PT Example adalah perusahaan teknologi terkemuka di Indonesia yang berfokus pada pengembangan solusi enterprise.",
        organization_context="Perusahaan ini memiliki pengalaman lebih dari 10 tahun dalam industri teknologi dan telah melayani berbagai klien korporat.",
        selected_tlos=[tlo1, tlo2],
        selected_performances=[perf1, perf2],
        selected_elos=[elo1, elo2, elo3],
        course_type="B2B"
    )
    
    # Create document generator
    generator = DocumentGenerator()
    
    # Generate document
    print("Generating syllabus document...")
    doc_bytes = generator.create_syllabus_document(materials)
    
    print(f"✓ Document generated successfully ({len(doc_bytes)} bytes)")
    
    # Verify it's a valid DOCX
    print("Verifying document structure...")
    doc = Document(BytesIO(doc_bytes))
    
    print(f"✓ Document has {len(doc.paragraphs)} paragraphs")
    
    # Extract and display content
    print("\nDocument content preview:")
    print("-" * 60)
    for i, para in enumerate(doc.paragraphs[:15]):  # Show first 15 paragraphs
        if para.text.strip():
            print(f"{i+1}. {para.text[:80]}{'...' if len(para.text) > 80 else ''}")
    
    # Save to file for manual inspection
    output_file = "test_syllabus_output.docx"
    with open(output_file, 'wb') as f:
        f.write(doc_bytes)
    
    print(f"\n✓ Document saved to {output_file}")
    print("\nAll tests passed! ✓")


if __name__ == "__main__":
    test_document_generator()
