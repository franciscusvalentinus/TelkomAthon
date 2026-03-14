"""
Manual test script for DocumentProcessor to verify implementation.
"""

from io import BytesIO
from docx import Document
from src.processors import DocumentProcessor, UnsupportedFormatError, EmptyFileError


def test_txt_extraction():
    """Test TXT file extraction."""
    print("Testing TXT extraction...")
    processor = DocumentProcessor()
    
    # Test valid TXT
    content = b"This is a test document.\nWith multiple lines."
    result = processor.extract_text_from_txt(content)
    print(f"✓ TXT extraction successful: {len(result)} characters")
    assert "test document" in result
    
    # Test empty TXT
    try:
        processor.extract_text_from_txt(b"")
        print("✗ Should have raised EmptyFileError")
    except EmptyFileError:
        print("✓ Empty TXT correctly raises EmptyFileError")


def test_docx_extraction():
    """Test DOCX file extraction."""
    print("\nTesting DOCX extraction...")
    processor = DocumentProcessor()
    
    # Create a sample DOCX
    doc = Document()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("With multiple paragraphs.")
    
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    result = processor.extract_text_from_docx(docx_file.read())
    print(f"✓ DOCX extraction successful: {len(result)} characters")
    assert "test DOCX document" in result


def test_process_document():
    """Test process_document dispatcher."""
    print("\nTesting process_document dispatcher...")
    processor = DocumentProcessor()
    
    # Test TXT
    txt_content = b"Sample text content for testing."
    result = processor.process_document(txt_content, '.txt')
    print(f"✓ TXT processing successful: {len(result)} characters")
    
    # Test case insensitivity
    result = processor.process_document(txt_content, 'TXT')
    print("✓ Case-insensitive file type handling works")
    
    # Test unsupported format
    try:
        processor.process_document(b"content", '.jpg')
        print("✗ Should have raised UnsupportedFormatError")
    except UnsupportedFormatError as e:
        print(f"✓ Unsupported format correctly raises error: {str(e)}")


def test_format_support():
    """Test format support checking."""
    print("\nTesting format support checking...")
    processor = DocumentProcessor()
    
    # Supported formats
    assert processor.is_format_supported('.pdf')
    assert processor.is_format_supported('.docx')
    assert processor.is_format_supported('.txt')
    assert processor.is_format_supported('pdf')  # Without dot
    print("✓ All supported formats recognized")
    
    # Unsupported formats
    assert not processor.is_format_supported('.jpg')
    assert not processor.is_format_supported('.png')
    assert not processor.is_format_supported('.exe')
    print("✓ Unsupported formats correctly rejected")


def test_content_validation():
    """Test content validation."""
    print("\nTesting content validation...")
    processor = DocumentProcessor()
    
    # Valid content
    assert processor.validate_document_content("This is valid content with enough text.")
    print("✓ Valid content passes validation")
    
    # Invalid content
    assert not processor.validate_document_content("")
    assert not processor.validate_document_content("   ")
    assert not processor.validate_document_content("Short")
    print("✓ Invalid content fails validation")


if __name__ == "__main__":
    print("=" * 60)
    print("DocumentProcessor Manual Test Suite")
    print("=" * 60)
    
    try:
        test_txt_extraction()
        test_docx_extraction()
        test_process_document()
        test_format_support()
        test_content_validation()
        
        print("\n" + "=" * 60)
        print("✓ All tests passed successfully!")
        print("=" * 60)
    except Exception as e:
        print(f"\n✗ Test failed with error: {e}")
        import traceback
        traceback.print_exc()
