"""
Document generator module for creating formatted syllabus documents.

This module provides functionality to generate DOCX syllabus documents
with consistent formatting and structure.
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from src.models.entities import SyllabusMaterials, TLO, Performance, ELO


class DocumentGenerator:
    """
    Handles generation of formatted DOCX syllabus documents.
    
    Creates professional syllabus documents with consistent styling
    including organization profile, TLOs, performances, and ELOs.
    """
    
    def __init__(self):
        """Initialize the document generator."""
        pass
    
    def create_syllabus_document(self, materials: SyllabusMaterials) -> bytes:
        """
        Generate DOCX syllabus document from selected materials.
        
        Args:
            materials: SyllabusMaterials containing all selected content
            
        Returns:
            DOCX document as bytes
        """
        doc = Document()
        
        # Set up document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('Silabus Kursus', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add course type
        course_type_para = doc.add_paragraph()
        course_type_run = course_type_para.add_run(f'Jenis Kursus: {materials.course_type}')
        course_type_run.bold = True
        course_type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add spacing
        
        # Add organization section
        self._format_organization_section(doc, materials)
        
        # Add TLO section
        self._format_tlo_section(doc, materials.selected_tlos)
        
        # Add performance section
        self._format_performance_section(doc, materials.selected_performances)
        
        # Add ELO section
        self._format_elo_section(doc, materials.selected_elos)
        
        # Save to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
    
    def _format_organization_section(self, doc: Document, materials: SyllabusMaterials) -> None:
        """
        Add organization profile section to document.
        
        Args:
            doc: The Document object to add content to
            materials: SyllabusMaterials containing organization info
        """
        # Add section heading
        doc.add_heading('Profil Organisasi', level=1)
        
        # Add summary
        doc.add_heading('Ringkasan', level=2)
        summary_para = doc.add_paragraph(materials.organization_summary)
        summary_para.paragraph_format.space_after = Pt(12)
        
        # Add context overview
        doc.add_heading('Konteks Organisasi', level=2)
        context_para = doc.add_paragraph(materials.organization_context)
        context_para.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph()  # Add spacing
    
    def _format_tlo_section(self, doc: Document, tlos: list[TLO]) -> None:
        """
        Add TLO section to document.
        
        Args:
            doc: The Document object to add content to
            tlos: List of selected TLO objects
        """
        # Add section heading
        doc.add_heading('Terminal Learning Objectives (TLO)', level=1)
        
        if not tlos:
            doc.add_paragraph('Tidak ada TLO yang dipilih.')
            return
        
        # Add each TLO as a numbered item
        for idx, tlo in enumerate(tlos, start=1):
            tlo_para = doc.add_paragraph(style='List Number')
            tlo_para.add_run(tlo.text)
            tlo_para.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()  # Add spacing
    
    def _format_performance_section(self, doc: Document, performances: list[Performance]) -> None:
        """
        Add performance section to document.
        
        Args:
            doc: The Document object to add content to
            performances: List of selected Performance objects
        """
        # Add section heading
        doc.add_heading('Performance Objectives', level=1)
        
        if not performances:
            doc.add_paragraph('Tidak ada performance objectives yang dipilih.')
            return
        
        # Add each performance as a numbered item
        for idx, performance in enumerate(performances, start=1):
            perf_para = doc.add_paragraph(style='List Number')
            perf_para.add_run(performance.text)
            perf_para.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()  # Add spacing
    
    def _format_elo_section(self, doc: Document, elos: list[ELO]) -> None:
        """
        Add ELO section to document.
        
        Args:
            doc: The Document object to add content to
            elos: List of selected ELO objects
        """
        # Add section heading
        doc.add_heading('Enabling Learning Objectives (ELO)', level=1)
        
        if not elos:
            doc.add_paragraph('Tidak ada ELO yang dipilih.')
            return
        
        # Add each ELO as a numbered item
        for idx, elo in enumerate(elos, start=1):
            elo_para = doc.add_paragraph(style='List Number')
            elo_para.add_run(elo.text)
            elo_para.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()  # Add spacing
