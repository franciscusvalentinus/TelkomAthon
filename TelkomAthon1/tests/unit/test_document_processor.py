"""
Unit tests for DocumentProcessor class.

Tests specific document format extraction and error handling scenarios.
"""

import pytest
from io import BytesIO
from docx import Document
import PyPDF2
from src.processors import DocumentProcessor, UnsupportedFormatError, EmptyFileError


@pytest.fixture
def processor():
    """Create a DocumentProcessor instance for testing."""
    return DocumentProcessor()


@pytest.fixture
def sample_txt_content():
    """Create sample TXT file content."""
    return b"This is a sample text document.\nIt has multiple lines.\nFor testing purposes."


@pytest.fixture
def sample_docx_content():
    """Create sample DOCX file content."""
    doc = Document()
    doc.add_paragraph("This is a sample DOCX document.")
    doc.add_paragraph("It has multiple paragraphs.")
    doc.add_paragraph("For testing purposes.")
    
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file.read()


@pytest.fixture
def sample_pdf_content():
    """Create sample PDF file content."""
    pdf_writer = PyPDF2.PdfWriter()
    
    # Create a simple PDF with text
    # Note: PyPDF2 doesn't easily create PDFs with text, so we'll use a minimal approach
    # In real tests, you'd use a pre-made PDF file
    pdf_file = BytesIO()
    pdf_writer.write(pdf_file)
    pdf_file.seek(0)
    return pdf_file.read()


class TestDocumentProcessor:
    """Test suite for DocumentProcessor class."""
    
    def test_extract_text_from_txt_utf8(self, processor, sample_txt_content):
        """Test extracting text from UTF-8 encoded TXT file."""
        result = processor.extract_text_from_txt(sample_txt_content)
        
        assert "sample text document" in result
        assert "multiple lines" in result
        assert len(result) > 0
    
    def test_extract_text_from_txt_with_bom(self, processor):
        """Test extracting text from TXT file with UTF-8 BOM."""
        content = b'\xef\xbb\xbfThis is UTF-8 with BOM'
        result = processor.extract_text_from_txt(content)
        
        assert "UTF-8 with BOM" in result
    
    def test_extract_text_from_txt_latin1(self, processor):
        """Test extracting text from Latin-1 encoded TXT file."""
        content = "Café résumé".encode('latin-1')
        result = processor.extract_text_from_txt(content)
        
        assert "Caf" in result or "r" in result  # Should decode successfully

    def test_extract_text_from_txt_empty_file(self, processor):
        """Test that empty TXT file raises EmptyFileError."""
        empty_content = b""
        
        with pytest.raises(EmptyFileError, match="TXT file is empty"):
            processor.extract_text_from_txt(empty_content)
    
    def test_extract_text_from_txt_whitespace_only(self, processor):
        """Test that TXT file with only whitespace raises EmptyFileError."""
        whitespace_content = b"   \n\t\n   "
        
        with pytest.raises(EmptyFileError, match="TXT file is empty"):
            processor.extract_text_from_txt(whitespace_content)
    
    def test_extract_text_from_docx(self, processor, sample_docx_content):
        """Test extracting text from DOCX file."""
        result = processor.extract_text_from_docx(sample_docx_content)
        
        assert "sample DOCX document" in result
        assert "multiple paragraphs" in result
        assert len(result) > 0
    
    def test_extract_text_from_docx_empty(self, processor):
        """Test that empty DOCX file raises EmptyFileError."""
        # Create an empty DOCX
        doc = Document()
        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        with pytest.raises(EmptyFileError, match="DOCX file contains no extractable text"):
            processor.extract_text_from_docx(docx_file.read())
    
    def test_extract_text_from_docx_corrupted(self, processor):
        """Test that corrupted DOCX file raises EmptyFileError."""
        corrupted_content = b"This is not a valid DOCX file"
        
        with pytest.raises(EmptyFileError):
            processor.extract_text_from_docx(corrupted_content)
    
    def test_process_document_txt(self, processor, sample_txt_content):
        """Test process_document with TXT file."""
        result = processor.process_document(sample_txt_content, '.txt')
        
        assert "sample text document" in result
        assert len(result) > 0
    
    def test_process_document_txt_without_dot(self, processor, sample_txt_content):
        """Test process_document with TXT file type without leading dot."""
        result = processor.process_document(sample_txt_content, 'txt')
        
        assert "sample text document" in result
        assert len(result) > 0
    
    def test_process_document_docx(self, processor, sample_docx_content):
        """Test process_document with DOCX file."""
        result = processor.process_document(sample_docx_content, '.docx')
        
        assert "sample DOCX document" in result
        assert len(result) > 0
    
    def test_process_document_unsupported_format(self, processor):
        """Test that unsupported file format raises UnsupportedFormatError."""
        content = b"Some content"
        
        with pytest.raises(UnsupportedFormatError, match="Unsupported file format"):
            processor.process_document(content, '.jpg')
    
    def test_process_document_case_insensitive(self, processor, sample_txt_content):
        """Test that file type is case-insensitive."""
        result = processor.process_document(sample_txt_content, '.TXT')
        
        assert "sample text document" in result

    def test_is_format_supported_pdf(self, processor):
        """Test that PDF format is recognized as supported."""
        assert processor.is_format_supported('.pdf') is True
        assert processor.is_format_supported('pdf') is True
        assert processor.is_format_supported('.PDF') is True
    
    def test_is_format_supported_docx(self, processor):
        """Test that DOCX format is recognized as supported."""
        assert processor.is_format_supported('.docx') is True
        assert processor.is_format_supported('docx') is True
    
    def test_is_format_supported_txt(self, processor):
        """Test that TXT format is recognized as supported."""
        assert processor.is_format_supported('.txt') is True
        assert processor.is_format_supported('txt') is True
    
    def test_is_format_supported_unsupported(self, processor):
        """Test that unsupported formats are recognized."""
        assert processor.is_format_supported('.jpg') is False
        assert processor.is_format_supported('.png') is False
        assert processor.is_format_supported('.exe') is False
        assert processor.is_format_supported('.doc') is False
    
    def test_validate_document_content_valid(self, processor):
        """Test validation of valid document content."""
        valid_text = "This is a valid document with meaningful content."
        assert processor.validate_document_content(valid_text) is True
    
    def test_validate_document_content_empty(self, processor):
        """Test validation of empty content."""
        assert processor.validate_document_content("") is False
        assert processor.validate_document_content(None) is False
    
    def test_validate_document_content_whitespace(self, processor):
        """Test validation of whitespace-only content."""
        assert processor.validate_document_content("   \n\t  ") is False
    
    def test_validate_document_content_too_short(self, processor):
        """Test validation of content that's too short to be meaningful."""
        assert processor.validate_document_content("Hi") is False
        assert processor.validate_document_content("Test") is False
    
    def test_validate_document_content_minimum_length(self, processor):
        """Test validation of content at minimum length threshold."""
        # Exactly 10 characters (minimum threshold)
        assert processor.validate_document_content("1234567890") is True
        # Just below threshold
        assert processor.validate_document_content("123456789") is False
    
    def test_supported_formats_constant(self, processor):
        """Test that SUPPORTED_FORMATS contains expected formats."""
        assert '.pdf' in processor.SUPPORTED_FORMATS
        assert '.docx' in processor.SUPPORTED_FORMATS
        assert '.txt' in processor.SUPPORTED_FORMATS
        assert len(processor.SUPPORTED_FORMATS) == 3
