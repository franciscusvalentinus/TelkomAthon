"""
Example usage of DocumentProcessor class.

This script demonstrates how to use the DocumentProcessor to extract text
from various document formats (PDF, DOCX, TXT).
"""

from io import BytesIO
from docx import Document
from src.processors import DocumentProcessor, UnsupportedFormatError, EmptyFileError


def example_txt_processing():
    """Example: Processing a TXT file."""
    print("=" * 60)
    print("Example 1: Processing TXT File")
    print("=" * 60)
    
    processor = DocumentProcessor()
    
    # Create sample TXT content
    txt_content = b"""This is a sample organization profile.

Our company specializes in educational technology and training solutions.
We provide innovative learning experiences for corporate clients.

Mission: To transform education through technology.
Vision: A world where learning is accessible to everyone."""
    
    try:
        # Extract text from TXT
        extracted_text = processor.extract_text_from_txt(txt_content)
        print(f"Extracted {len(extracted_text)} characters from TXT file")
        print("\nExtracted content:")
        print("-" * 60)
        print(extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text)
        print("-" * 60)
        
        # Validate content
        is_valid = processor.validate_document_content(extracted_text)
        print(f"\nContent validation: {'✓ Valid' if is_valid else '✗ Invalid'}")
        
    except EmptyFileError as e:
        print(f"Error: {e}")


def example_docx_processing():
    """Example: Processing a DOCX file."""
    print("\n" + "=" * 60)
    print("Example 2: Processing DOCX File")
    print("=" * 60)
    
    processor = DocumentProcessor()
    
    # Create sample DOCX content
    doc = Document()
    doc.add_heading('Organization Profile', 0)
    doc.add_paragraph('Company Name: TechEdu Solutions')
    doc.add_paragraph('Industry: Educational Technology')
    doc.add_paragraph('Founded: 2020')
    doc.add_paragraph('')
    doc.add_heading('About Us', 1)
    doc.add_paragraph(
        'We are a leading provider of AI-powered educational solutions. '
        'Our platform helps organizations create engaging learning experiences.'
    )
    
    # Save to BytesIO
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    docx_content = docx_file.read()
    
    try:
        # Extract text from DOCX
        extracted_text = processor.extract_text_from_docx(docx_content)
        print(f"Extracted {len(extracted_text)} characters from DOCX file")
        print("\nExtracted content:")
        print("-" * 60)
        print(extracted_text)
        print("-" * 60)
        
    except EmptyFileError as e:
        print(f"Error: {e}")


def example_process_document_dispatcher():
    """Example: Using the process_document dispatcher method."""
    print("\n" + "=" * 60)
    print("Example 3: Using process_document Dispatcher")
    print("=" * 60)
    
    processor = DocumentProcessor()
    
    # Sample content
    content = b"Sample organization profile for testing the dispatcher method."
    
    # Process with different file types
    file_types = ['.txt', 'TXT', '.docx', '.pdf']
    
    for file_type in file_types:
        try:
            if file_type.lower() in ['.txt', 'txt']:
                result = processor.process_document(content, file_type)
                print(f"✓ Successfully processed {file_type}: {len(result)} characters")
            else:
                # For demo purposes, we'll just check format support
                is_supported = processor.is_format_supported(file_type)
                print(f"{'✓' if is_supported else '✗'} Format {file_type}: "
                      f"{'Supported' if is_supported else 'Not supported'}")
        except UnsupportedFormatError as e:
            print(f"✗ {file_type}: {e}")


def example_error_handling():
    """Example: Error handling for various scenarios."""
    print("\n" + "=" * 60)
    print("Example 4: Error Handling")
    print("=" * 60)
    
    processor = DocumentProcessor()
    
    # Test 1: Empty file
    print("\n1. Testing empty file:")
    try:
        processor.extract_text_from_txt(b"")
    except EmptyFileError as e:
        print(f"   ✓ Caught expected error: {e}")
    
    # Test 2: Unsupported format
    print("\n2. Testing unsupported format:")
    try:
        processor.process_document(b"content", '.jpg')
    except UnsupportedFormatError as e:
        print(f"   ✓ Caught expected error: {e}")
    
    # Test 3: Whitespace-only content
    print("\n3. Testing whitespace-only content:")
    try:
        processor.extract_text_from_txt(b"   \n\t\n   ")
    except EmptyFileError as e:
        print(f"   ✓ Caught expected error: {e}")
    
    # Test 4: Content validation
    print("\n4. Testing content validation:")
    test_cases = [
        ("Valid content", "This is a valid document with enough text.", True),
        ("Empty string", "", False),
        ("Too short", "Hi", False),
        ("Whitespace only", "   \n\t  ", False),
    ]
    
    for name, content, expected in test_cases:
        result = processor.validate_document_content(content)
        status = "✓" if result == expected else "✗"
        print(f"   {status} {name}: {result} (expected {expected})")


def example_format_checking():
    """Example: Checking supported formats."""
    print("\n" + "=" * 60)
    print("Example 5: Format Support Checking")
    print("=" * 60)
    
    processor = DocumentProcessor()
    
    print("\nSupported formats:")
    for fmt in processor.SUPPORTED_FORMATS:
        print(f"  ✓ {fmt}")
    
    print("\nTesting various format checks:")
    test_formats = [
        '.pdf', 'pdf', '.PDF',
        '.docx', 'DOCX',
        '.txt', 'txt',
        '.jpg', '.png', '.doc', '.xlsx'
    ]
    
    for fmt in test_formats:
        is_supported = processor.is_format_supported(fmt)
        status = "✓" if is_supported else "✗"
        print(f"  {status} {fmt}: {'Supported' if is_supported else 'Not supported'}")


if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("DocumentProcessor Usage Examples")
    print("=" * 60)
    
    example_txt_processing()
    example_docx_processing()
    example_process_document_dispatcher()
    example_error_handling()
    example_format_checking()
    
    print("\n" + "=" * 60)
    print("Examples completed!")
    print("=" * 60)
